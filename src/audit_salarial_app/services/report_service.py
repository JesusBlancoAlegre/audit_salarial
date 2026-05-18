import os
from datetime import datetime
from flask import current_app
from docx import Document
from docx.shared import Pt, Inches
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors

from ..extensions import db
from ..models import Auditoria, AuditoriaArchivo, Resultado, Dimension

def generar_informe_word(auditoria_id):
    auditoria = db.session.get(Auditoria, auditoria_id)
    if not auditoria:
        return False, "Auditoría no encontrada"
        
    doc = Document()
    doc.add_heading(f"Informe Técnico de Auditoría Salarial #{auditoria_id}", 0)
    
    p = doc.add_paragraph(f"Empresa: {auditoria.empresa.nombre}\n")
    p.add_run(f"Fecha de generación: {datetime.now().strftime('%Y-%m-%d %H:%M')}\n")
    
    # Global
    doc.add_heading("1. Análisis Global", level=1)
    res_global = Resultado.query.join(Dimension).filter(
        Resultado.auditoria_id == auditoria_id,
        Dimension.codigo == 'GLOBAL'
    ).first()
    
    if res_global:
        doc.add_paragraph(f"Nº Total Empleados: {res_global.n_total}")
        doc.add_paragraph(f"Hombres: {res_global.n_hombres} - Mujeres: {res_global.n_mujeres}")
        doc.add_paragraph(f"Brecha Media: {res_global.brecha_media_pct:.2f}%")
        doc.add_paragraph(f"Brecha Mediana: {res_global.brecha_mediana_pct:.2f}%")
    else:
        doc.add_paragraph("No hay datos globales calculados.")
        
    # Grupos
    doc.add_heading("2. Análisis por Grupo Profesional", level=1)
    res_grupos = Resultado.query.join(Dimension).filter(
        Resultado.auditoria_id == auditoria_id,
        Dimension.codigo == 'GRUPO_PROFESIONAL'
    ).order_by(Resultado.dimension_valor).all()
    
    if res_grupos:
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Grupo'
        hdr_cells[1].text = 'Media Hombres'
        hdr_cells[2].text = 'Media Mujeres'
        hdr_cells[3].text = 'Brecha (%)'
        
        for res in res_grupos:
            row_cells = table.add_row().cells
            row_cells[0].text = str(res.dimension_valor)
            row_cells[1].text = f"{res.media_hombres:.2f} €" if res.media_hombres else "0.00 €"
            row_cells[2].text = f"{res.media_mujeres:.2f} €" if res.media_mujeres else "0.00 €"
            row_cells[3].text = f"{res.brecha_media_pct:.2f}%" if res.brecha_media_pct else "0.00%"
    else:
        doc.add_paragraph("No hay datos por grupos calculados.")
        
    filename = f"Informe_Tecnico_Audit_{auditoria_id}_{datetime.now().strftime('%Y%m%d%H%M')}.docx"
    filepath = os.path.join(current_app.config['UPLOAD_FOLDER'], filename)
    doc.save(filepath)
    
    aa = AuditoriaArchivo(
        auditoria_id=auditoria_id,
        tipo='WORD_TECNICO',
        ruta=filepath,
        nombre=filename
    )
    db.session.add(aa)
    db.session.commit()
    
    return True, filename

def generar_informe_pdf(auditoria_id):
    auditoria = db.session.get(Auditoria, auditoria_id)
    if not auditoria:
        return False, "Auditoría no encontrada"
        
    filename = f"Informe_Ejecutivo_Audit_{auditoria_id}_{datetime.now().strftime('%Y%m%d%H%M')}.pdf"
    filepath = os.path.join(current_app.config['UPLOAD_FOLDER'], filename)
    
    doc = SimpleDocTemplate(filepath, pagesize=letter)
    styles = getSampleStyleSheet()
    elements = []
    
    elements.append(Paragraph(f"Informe Ejecutivo - Auditoría Salarial #{auditoria_id}", styles['Title']))
    elements.append(Spacer(1, 12))
    
    elements.append(Paragraph(f"<b>Empresa:</b> {auditoria.empresa.nombre}", styles['Normal']))
    elements.append(Paragraph(f"<b>Fecha:</b> {datetime.now().strftime('%Y-%m-%d')}", styles['Normal']))
    elements.append(Spacer(1, 24))
    
    res_global = Resultado.query.join(Dimension).filter(
        Resultado.auditoria_id == auditoria_id,
        Dimension.codigo == 'GLOBAL'
    ).first()
    
    if res_global:
        elements.append(Paragraph("Resumen Global", styles['Heading2']))
        
        data = [
            ["Métrica", "Valor"],
            ["Total Empleados", str(res_global.n_total)],
            ["Hombres", str(res_global.n_hombres)],
            ["Mujeres", str(res_global.n_mujeres)],
            ["Brecha Salarial Media", f"{res_global.brecha_media_pct:.2f}%"],
            ["Brecha Salarial Mediana", f"{res_global.brecha_mediana_pct:.2f}%"]
        ]
        
        t = Table(data, colWidths=[200, 100])
        t.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ]))
        elements.append(t)
    
    doc.build(elements)
    
    aa = AuditoriaArchivo(
        auditoria_id=auditoria_id,
        tipo='PDF_EJECUTIVO',
        ruta=filepath,
        nombre=filename
    )
    db.session.add(aa)
    db.session.commit()
    
    return True, filename

def generar_informe_individual_pdf(auditoria_id, anomalia_id):
    from audit_salarial_app.models import Anomalia, Resultado
    anomalia = db.session.get(Anomalia, anomalia_id)
    auditoria = db.session.get(Auditoria, auditoria_id)
    if not anomalia or not auditoria:
        return False, "Datos no encontrados"
        
    resultado_grupo = Resultado.query.filter_by(
        auditoria_id=auditoria_id,
        dimension_valor=anomalia.dimension_valor
    ).first()
    
    filename = f"Informe_Individual_Fila_{anomalia.id_fila_excel or 'ND'}_{datetime.now().strftime('%Y%m%d%H%M')}.pdf"
    filepath = os.path.join(current_app.config['UPLOAD_FOLDER'], filename)
    
    doc = SimpleDocTemplate(filepath, pagesize=letter)
    styles = getSampleStyleSheet()
    elements = []
    
    elements.append(Paragraph(f"Informe Individual de Transparencia Retributiva (RD 902/2020)", styles['Title']))
    elements.append(Spacer(1, 12))
    
    elements.append(Paragraph(f"<b>Empresa:</b> {auditoria.empresa.nombre}", styles['Normal']))
    elements.append(Paragraph(f"<b>Auditoría ID:</b> #{auditoria_id}", styles['Normal']))
    elements.append(Paragraph(f"<b>Fecha de Informe:</b> {datetime.now().strftime('%Y-%m-%d')}", styles['Normal']))
    elements.append(Spacer(1, 24))
    
    # 1. Datos del Empleado
    elements.append(Paragraph("1. Datos de Identificación", styles['Heading2']))
    data_emp = [
        ["Fila Excel Referencia", str(anomalia.id_fila_excel or 'N/D')],
        ["Grupo Profesional", str(anomalia.dimension_valor)],
        ["Retribución Detectada", f"{anomalia.valor:,.2f} €"]
    ]
    t1 = Table(data_emp, colWidths=[200, 200])
    t1.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (0, -1), colors.lightgrey),
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ('PADDING', (0, 0), (-1, -1), 6),
    ]))
    elements.append(t1)
    elements.append(Spacer(1, 12))
    
    # 2. Análisis de Anomalía
    elements.append(Paragraph("2. Análisis de Atipicidad (Desviación)", styles['Heading2']))
    elements.append(Paragraph("Este salario ha sido marcado como un valor atípico estadístico dentro de su grupo.", styles['Normal']))
    elements.append(Spacer(1, 6))
    
    data_anom = [
        ["Método de Detección", str(anomalia.metodo)],
        ["Severidad", str(anomalia.severidad)],
        ["Umbral Superior Grupo", f"{anomalia.umbral_superior:,.2f} €" if anomalia.umbral_superior else "N/D"],
        ["Umbral Inferior Grupo", f"{anomalia.umbral_inferior:,.2f} €" if anomalia.umbral_inferior else "N/D"],
        ["Z-Score (Desviación)", f"{anomalia.z_score:,.2f}" if anomalia.z_score else "N/D"]
    ]
    t2 = Table(data_anom, colWidths=[200, 200])
    t2.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (0, -1), colors.lightgrey),
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ('PADDING', (0, 0), (-1, -1), 6),
    ]))
    elements.append(t2)
    elements.append(Spacer(1, 12))
    
    # 3. Contexto del Grupo
    elements.append(Paragraph("3. Contexto del Grupo Profesional (Igualdad Retributiva)", styles['Heading2']))
    if resultado_grupo:
        elements.append(Paragraph(f"La brecha salarial actual en este grupo profesional es del <b>{resultado_grupo.brecha_media_pct:.2f}%</b>.", styles['Normal']))
        elements.append(Spacer(1, 6))
        data_grupo = [
            ["Salario Medio Hombres", f"{resultado_grupo.media_hombres:,.2f} €" if resultado_grupo.media_hombres else "N/D"],
            ["Salario Medio Mujeres", f"{resultado_grupo.media_mujeres:,.2f} €" if resultado_grupo.media_mujeres else "N/D"],
            ["Total Empleados en Grupo", str(resultado_grupo.n_total)]
        ]
        t3 = Table(data_grupo, colWidths=[200, 200])
        t3.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (0, -1), colors.lightgrey),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ('PADDING', (0, 0), (-1, -1), 6),
        ]))
        elements.append(t3)
    else:
        elements.append(Paragraph("No hay datos comparativos del grupo disponibles.", styles['Normal']))
        
    elements.append(Spacer(1, 12))
    
    # 4. Obligaciones Legales
    elements.append(Paragraph("4. Justificación Objetiva (Art. 28 ET y RD 902/2020)", styles['Heading2']))
    elements.append(Paragraph("<b>Nota Legal:</b> Dado que este salario representa una anomalía estadística frente a la retribución media de trabajos de igual valor, la empresa debe justificar de manera objetiva, razonable y transparente los factores determinantes de esta diferencia (por ejemplo: pluses de antigüedad, turnicidad, especial calificación, etc.). De lo contrario, podría considerarse un indicio de discriminación retributiva.", styles['Normal']))
    
    doc.build(elements)
    
    return True, filename

def generar_plan_actuacion_pdf(auditoria_id):
    """Genera un PDF formal del Plan de Actuación (RD 902/2020)."""
    from ..models import PlanActuacion, AuditoriaRecomendacion, Tarea

    auditoria = db.session.get(Auditoria, auditoria_id)
    if not auditoria:
        return False, "Auditoría no encontrada"

    plan = PlanActuacion.query.filter_by(auditoria_id=auditoria_id).first()
    res_global = Resultado.query.join(Dimension).filter(
        Resultado.auditoria_id == auditoria_id, Dimension.codigo == 'GLOBAL'
    ).first()
    recomendaciones = AuditoriaRecomendacion.query.filter_by(auditoria_id=auditoria_id).all()
    tareas = Tarea.query.filter_by(auditoria_id=auditoria_id).order_by(Tarea.fecha_inicio).all()

    filename = f"Plan_Actuacion_Audit_{auditoria_id}_{datetime.now().strftime('%Y%m%d%H%M')}.pdf"
    filepath = os.path.join(current_app.config['UPLOAD_FOLDER'], filename)

    doc = SimpleDocTemplate(filepath, pagesize=letter)
    styles = getSampleStyleSheet()
    elements = []

    # Title
    elements.append(Paragraph("Plan de Actuación para la Igualdad Retributiva", styles['Title']))
    elements.append(Paragraph(f"(Conforme al Real Decreto 902/2020)", styles['Normal']))
    elements.append(Spacer(1, 12))
    elements.append(Paragraph(f"<b>Empresa:</b> {auditoria.empresa.nombre}", styles['Normal']))
    elements.append(Paragraph(f"<b>Auditoría ID:</b> #{auditoria_id}", styles['Normal']))
    elements.append(Paragraph(f"<b>Fecha:</b> {datetime.now().strftime('%d/%m/%Y')}", styles['Normal']))
    elements.append(Spacer(1, 24))

    # Section 1: Diagnostic
    elements.append(Paragraph("1. Diagnóstico de la Situación Retributiva", styles['Heading2']))
    if res_global:
        diag_data = [
            ["Métrica", "Valor"],
            ["Total Empleados", str(res_global.n_total)],
            ["Hombres", str(res_global.n_hombres)],
            ["Mujeres", str(res_global.n_mujeres)],
            ["Brecha Salarial Media", f"{res_global.brecha_media_pct:.2f}%"],
            ["Brecha Salarial Mediana", f"{res_global.brecha_mediana_pct:.2f}%"],
            ["Nivel de Riesgo", res_global.nivel_riesgo or "N/D"],
            ["Score de Riesgo", f"{res_global.score_riesgo:.1f}/100" if res_global.score_riesgo else "N/D"],
        ]
        t = Table(diag_data, colWidths=[220, 120])
        t.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#4338ca')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor('#f0f0ff')),
            ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#c7d2fe')),
            ('ALIGN', (1, 0), (1, -1), 'CENTER'),
        ]))
        elements.append(t)
    else:
        elements.append(Paragraph("No hay datos de diagnóstico calculados.", styles['Normal']))
    elements.append(Spacer(1, 18))

    # Section 2: Objectives
    elements.append(Paragraph("2. Objetivos del Plan", styles['Heading2']))
    if plan:
        elements.append(Paragraph(f"<b>Objetivo General:</b> {plan.objetivo_general or 'No definido'}", styles['Normal']))
        elements.append(Paragraph(f"<b>Objetivo de Brecha:</b> {plan.objetivo_brecha_pct}%" if plan.objetivo_brecha_pct else "<b>Objetivo de Brecha:</b> No definido", styles['Normal']))
        elements.append(Paragraph(f"<b>Plazo:</b> {plan.plazo_meses} meses", styles['Normal']))
        elements.append(Paragraph(f"<b>Estado:</b> {plan.estado}", styles['Normal']))
    else:
        elements.append(Paragraph("Plan de actuación no definido aún.", styles['Normal']))
    elements.append(Spacer(1, 18))

    # Section 3: Corrective Measures
    elements.append(Paragraph("3. Medidas Correctoras", styles['Heading2']))
    if recomendaciones:
        rec_data = [["Prioridad", "Medida", "Impacto", "Plazo", "Estado"]]
        for r in recomendaciones:
            rec_data.append([
                r.prioridad,
                r.recomendacion.titulo if r.recomendacion else "N/D",
                f"{r.impacto_reduccion_pct}%" if r.impacto_reduccion_pct else "N/D",
                f"{r.meses_estimados} meses" if r.meses_estimados else "N/D",
                "Aplicada" if r.aplicada else "Pendiente",
            ])
        t2 = Table(rec_data, colWidths=[65, 170, 55, 55, 55])
        t2.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#b45309')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 8),
            ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor('#fef3c7')),
            ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#fbbf24')),
        ]))
        elements.append(t2)
    else:
        elements.append(Paragraph("No se han generado medidas correctoras.", styles['Normal']))
    elements.append(Spacer(1, 18))

    # Section 4: Timeline
    elements.append(Paragraph("4. Cronograma de Actuación", styles['Heading2']))
    if tareas:
        tar_data = [["Tarea", "Tipo", "Inicio", "Fin", "Estado"]]
        for t_item in tareas:
            tar_data.append([
                t_item.titulo,
                t_item.tipo,
                t_item.fecha_inicio.strftime('%d/%m/%Y'),
                t_item.fecha_fin.strftime('%d/%m/%Y'),
                t_item.estado,
            ])
        t3 = Table(tar_data, colWidths=[150, 60, 70, 70, 70])
        t3.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#059669')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 8),
            ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor('#ecfdf5')),
            ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#6ee7b7')),
        ]))
        elements.append(t3)
    else:
        elements.append(Paragraph("No hay tareas asignadas al cronograma.", styles['Normal']))
    elements.append(Spacer(1, 18))

    # Section 5: Legal
    elements.append(Paragraph("5. Marco Legal Aplicable", styles['Heading2']))
    elements.append(Paragraph(
        "<b>Real Decreto 902/2020</b>, de 13 de octubre, de igualdad retributiva entre mujeres y hombres. "
        "El presente plan de actuación se elabora conforme a los artículos 6 a 9 del citado Real Decreto, "
        "que establece la obligación de las empresas de llevar a cabo una auditoría retributiva y de "
        "establecer un plan de actuación con objetivos concretos, medidas correctoras y un cronograma "
        "de implementación para garantizar la igualdad retributiva efectiva.",
        styles['Normal']
    ))

    doc.build(elements)
    return True, filename
